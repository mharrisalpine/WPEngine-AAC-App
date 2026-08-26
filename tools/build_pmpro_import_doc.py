from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "aac-pmpro-member-import-instructions.md"
OUTPUT = ROOT / "docs" / "AAC PMPro Member Import Instructions.docx"

CONTENT_WIDTH_DXA = 9360
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "AAB7C4"
BODY = "1F2933"


def set_cell_text(cell, text, bold=False, color=BODY, size=9.5, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Consolas" if mono else "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    if cols == 2:
        widths = [2700, CONTENT_WIDTH_DXA - 2700]
    elif cols == 3:
        widths = [2500, 3430, 3430]
    elif cols == 4:
        widths = [2600, 1900, 2300, 2560]
    elif cols == 5:
        widths = [2700, 1600, 1300, 1700, 2060]
    elif cols == 6:
        widths = [2500, 1450, 1350, 1350, 1450, 1260]
    else:
        base = CONTENT_WIDTH_DXA // cols
        widths = [base] * cols
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)

    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, text.strip(), bold=row_index == 0, color="0B2545" if row_index == 0 else BODY)
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)

    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("111827")
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph()


def split_table_row(line):
    stripped = line.strip().strip("|")
    return [part.strip().strip("`") for part in stripped.split("|")]


def is_separator_row(line):
    return bool(re.fullmatch(r"\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*", line))


def clean_inline(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def add_paragraph_with_inline_code(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6 if style is None else 4)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string("111827")
        else:
            run = p.add_run(part.replace("**", ""))
            run.font.color.rgb = RGBColor.from_string(BODY)
    return p


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title_block(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("AAC PMPro Member Import Instructions")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Stripe subscription preservation, PMPro membership mapping, family-plan imports, donations, and Salesforce reporting.")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("53616F")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run(f"Prepared {date.today().strftime('%B %-d, %Y')}")
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("6B7280")


def build_document():
    doc = Document()
    setup_styles(doc)
    add_title_block(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("# "):
            if stripped[2:].strip() != "AAC PMPro Member Import Instructions":
                doc.add_heading(clean_inline(stripped[2:].strip()), level=1)
            idx += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:].strip()), level=1)
            idx += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:].strip()), level=2)
            idx += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code_block(doc, code_lines)
            idx += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                if not is_separator_row(lines[idx]):
                    table_lines.append(split_table_row(lines[idx]))
                idx += 1
            add_table(doc, table_lines)
            continue

        if stripped.startswith("- "):
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                p = add_paragraph_with_inline_code(doc, lines[idx].strip()[2:], style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                idx += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            while idx < len(lines):
                numbered_match = re.match(r"^\d+\.\s+(.*)$", lines[idx].strip())
                if not numbered_match:
                    break
                p = add_paragraph_with_inline_code(doc, numbered_match.group(1), style="List Number")
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                idx += 1
            continue

        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            next_line = lines[idx].strip()
            if not next_line or next_line.startswith(("#", "```", "|", "- ")) or re.match(r"^\d+\.\s+", next_line):
                break
            paragraph_lines.append(next_line)
            idx += 1
        add_paragraph_with_inline_code(doc, " ".join(paragraph_lines))

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("AAC PMPro Import Instructions")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
